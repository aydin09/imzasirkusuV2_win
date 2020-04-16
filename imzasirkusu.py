from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.shared import Cm, Inches
from docx.shared import Length
import os
from tkinter import *
import sqlite3
import tkinter.ttk as ttk
import locale
locale.setlocale(locale.LC_ALL, "")

def bilgi_girişi(event):
    liste1=liste.get(ACTIVE)

    mudur_yard_adi.delete(0,END)
    personel_adi_soyadi.delete(0,END)
    personel_gorev_brans.delete(0,END)
                       
    vt = sqlite3.connect(str(liste1)+'.sq3')
    im= vt.cursor()
    im.execute(""" SELECT * FROM imza""")
    rows = im.fetchall()
    data_str = ""
    sf = "{}{}"
    for row in rows:
        data_str += sf.format(row[0], row[1])

        personel_adi_soyadi.insert(END,row[0])
        personel_gorev_brans.insert(END,row[1])

def bilgi_girişi2(event):
    liste22=liste2.get(ACTIVE)

    mudur_yard_adi.delete(0,END)
    personel_adi_soyadi.delete(0,END)
    personel_gorev_brans.delete(0,END)
                       
    vt = sqlite3.connect(str(liste22)+'.sql3')
    im= vt.cursor()
    im.execute("""SELECT * FROM imza""")
    rows = im.fetchall()
    data_str = ""
    sf = "{}"
    for row in rows:
        data_str += sf.format(row[0])

        mudur_yard_adi.insert(END,row[0])
                        
def kaydet_personel():
    kaymakamlik1 = kaymakamlik.get()
    okul_adi1 = okul_adi.get()
    personel_adi_soyadi1 = personel_adi_soyadi.get()
    personel_gorev_brans1 = personel_gorev_brans.get()

    if personel_adi_soyadi1 =="" or personel_gorev_brans1 =="":
        uyari=Toplevel()
        Label(uyari, text ='Personel bilgilerini tam giriniz!').pack()

    else:    
        personel_adi_soyadi.delete(0,END)
        personel_gorev_brans.delete(0,END)

        vt1 = sqlite3.connect('kaymakamlik.sq')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS imza(kaymakamlik TEXT)""")
        im1.execute("""UPDATE imza SET  kaymakamlik=?""",(kaymakamlik1,))
        vt1.commit()        
        
        vt2 = sqlite3.connect('okuladi.sql')
        im2= vt2.cursor()
        im2.execute("""CREATE TABLE IF NOT EXISTS imza(okuladi TEXT)""")
        im2.execute("""UPDATE imza SET  okuladi=?""",(okul_adi1,))
        vt2.commit()

        if os.path.exists(personel_adi_soyadi1+'.sq3')== False:
            vt1 = sqlite3.connect(personel_adi_soyadi1+'.sq3')
            im1= vt1.cursor()
            im1.execute("""CREATE TABLE IF NOT EXISTS imza(personeladisoyadi TEXT, personelgorevbrans TEXT)""")
            im1.execute("""INSERT INTO imza VALUES  (?,?)""",(personel_adi_soyadi1, personel_gorev_brans1,))
            vt1.commit()

            liste.delete(0,END)

            for i in sorted(os.listdir(),key=locale.strxfrm):
                if i.endswith('.sq3'):
                    liste.insert(END,i[0:-4])

        else:
            vt2 = sqlite3.connect(personel_adi_soyadi1+'.sq3')
            im2= vt2.cursor()
            im2.execute("""CREATE TABLE IF NOT EXISTS imza(personeladisoyadi TEXT, personelgorevbrans TEXT)""")
            im2.execute("""UPDATE imza SET  personeladisoyadi=?, personelgorevbrans=?""",(personel_adi_soyadi1,personel_gorev_brans1,))
            
            vt2.commit()

def kaydet_müd_yard():
    kaymakamlik1 = kaymakamlik.get()
    okul_adi1 = okul_adi.get()
    mudur_yard_adi1 = mudur_yard_adi.get()

    if mudur_yard_adi1 =="":
        uyari=Toplevel()
        Label(uyari, text ='Müdür Yardımcısının adını soyadını giriniz!').pack()

    else:
        mudur_yard_adi.delete(0,END)
        personel_adi_soyadi.delete(0,END)
        personel_gorev_brans.delete(0,END)

        vt1 = sqlite3.connect('kaymakamlik.sq')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS imza(kaymakamlik TEXT)""")
        im1.execute("""UPDATE imza SET  kaymakamlik=?""",(kaymakamlik1,))
        vt1.commit()        
        
        vt2 = sqlite3.connect('okuladi.sql')
        im2= vt2.cursor()
        im2.execute("""CREATE TABLE IF NOT EXISTS imza(okuladi TEXT)""")
        im2.execute("""UPDATE imza SET  okuladi=?""",(okul_adi1,))
        vt2.commit()
        
        if os.path.exists(mudur_yard_adi1+'.sql3')== False:
            vt1 = sqlite3.connect(mudur_yard_adi1+'.sql3')
            im1= vt1.cursor()
            im1.execute("""CREATE TABLE IF NOT EXISTS imza(mudur_yard_adi TEXT)""")
            im1.execute("""INSERT INTO imza VALUES  (?)""",(mudur_yard_adi1,))
            vt1.commit()

            liste2.delete(0,END)

            for i in sorted(os.listdir(),key=locale.strxfrm):
                if i.endswith('.sql3'):
                    liste2.insert(END,i[0:-5])

        else:
            vt2 = sqlite3.connect(mudur_yard_adi1+'.sql3')
            im2= vt2.cursor()
            im2.execute("""CREATE TABLE IF NOT EXISTS imza(mudur_yard_adi TEXT)""")
            im2.execute("""UPDATE imza SET  mudur_yard_adi=?""",(mudur_yard_adi1,))
            
            vt2.commit()
       
def cikti():
    vt1 = sqlite3.connect('kaymakamlik.sq')
    im1= vt1.cursor()
    im1.execute("""CREATE TABLE IF NOT EXISTS imza(kaymakamlik TEXT)""")
    im1.execute("""SELECT * FROM  imza""")
    rows = im1.fetchall()
    data_str = ""
    sf = "{}"
    for rowkay in rows:
        data_str += sf.format(rowkay[0])
    vt1.commit()        
    
    vt2 = sqlite3.connect('okuladi.sql')
    im2= vt2.cursor()
    im2.execute("""CREATE TABLE IF NOT EXISTS imza(okuladi TEXT)""")
    im2.execute("""SELECT * FROM  imza""")
    rows = im2.fetchall()
    data_str1 = ""
    sf = "{}"
    for rowokul in rows:
        data_str1 += sf.format(rowokul[0])
    vt2.commit()

    personel=[]
    
    for i in sorted(os.listdir(),key=locale.strxfrm):
        if i.endswith('.sq3'):
            personel.append(i[0:-4])

    data3 = []
 
    while True:
        for s in range(0,len(personel)):
            
            vt3= sqlite3.connect(personel[s]+'.sq3')
            im3= vt3.cursor()
            im3.execute(""" SELECT * FROM imza""")
            rows3 = im3.fetchall()
            
            for row3 in rows3:
                data3.append(row3)

        break

    mud_yard=[]
    for i in sorted(os.listdir(),key=locale.strxfrm):
        if i.endswith('.sql3'):
            mud_yard.append(i[0:-5])

    data4 = []
 
    while True:
        for s in range(0,len(mud_yard)):
            
            vt4= sqlite3.connect(mud_yard[s]+'.sql3')
            im4= vt4.cursor()
            im4.execute("""SELECT * FROM imza""")
            rows4 = im4.fetchall()
            
            for row4 in rows4:
                data4.append(row4)

        break
    
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)    

    paragraph = document.add_paragraph()
    paragraph.add_run("T.C.").bold = True
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.alignment = 1
   
    paragraph = document.add_paragraph()
    paragraph.add_run(rowkay[0]).bold = True
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.alignment = 1

    paragraph = document.add_paragraph()
    paragraph.add_run(rowokul[0]).bold = True
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.alignment = 1

    paragraph = document.add_paragraph()
    paragraph.add_run("İMZA SİRKÜSÜ").bold = True
    paragraph.alignment = 1
             
    table = document.add_table(rows=len(personel)+len(mud_yard)+1, cols=4,style = 'Table Grid')
    
    cell = table.cell(0,0)
    table.cell(0,0).paragraphs[0].add_run("SIRA NO").bold = True
    table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(0.6)
    
    
    cell = table.cell(0,1)
    table.cell(0,1).paragraphs[0].add_run("PERSONELİN ADI SOYADI").bold = True
    table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[1].width = Inches(2.5)

    cell = table.cell(0,2)
    table.cell(0,2).paragraphs[0].add_run("GÖREVİ/BRANŞI").bold = True
    table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[2].width = Inches(2.5)

    cell = table.cell(0,3)
    table.cell(0,3).paragraphs[0].add_run("İMZA").bold = True
    table.cell(0,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[3].width = Inches(1.0)

    for s in range(0,len(mud_yard)):
        cell = table.cell(s+1,1)
        cell.text =data4[s][0]
        table.cell(s+1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        cell = table.cell(s+1,2)
        cell.text ="MÜDÜR YARDIMCISI"
        table.cell(s+1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT

        cell = table.cell(s+1,0)
        cell.text =str(int(s)+1)
        table.cell(s+1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

    for s1 in range(0,len(personel)):
        cell = table.cell(s1+len(mud_yard)+1,1)
        cell.text =data3[s1][0]
        table.cell(len(mud_yard)+1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        cell = table.cell(s1+len(mud_yard)+1,2)
        cell.text =data3[s1][1]
        table.cell(len(mud_yard)+1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT

        cell = table.cell(s1+len(mud_yard)+1,0)
        cell.text =str(int(s1)+len(mud_yard)+1)
        table.cell(s1+len(mud_yard)+1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            
    document.save('imzasirkusu.docx')

    os.startfile("imzasirkusu.docx")

def sil_personel():
    data_sil=liste.get(ACTIVE)

    os.remove(data_sil+".sq3")

    liste.delete(0,END)

    for i in sorted(os.listdir(os.getcwd())):
        if i.endswith('.sq3'):
            liste.insert(END,i[0:-4])

    mudur_yard_adi.delete(0,END)
    personel_adi_soyadi.delete(0,END)
    personel_gorev_brans.delete(0,END)

def sil_müd_yard():
    data_sil1=liste2.get(ACTIVE)

    os.remove(data_sil1+".sql3")

    liste2.delete(0,END)

    for i in sorted(os.listdir(os.getcwd())):
        if i.endswith('.sql3'):
            liste2.insert(END,i[0:-5])

    mudur_yard_adi.delete(0,END)
    personel_adi_soyadi.delete(0,END)
    personel_gorev_brans.delete(0,END)

root = Tk()
root.title("İmza Sirküsü")
root.resizable(width=FALSE ,height=FALSE)
img=PhotoImage(file='imza.png')
root.tk.call('wm','iconphoto',root._w,img)
mainframe = ttk.Frame(root,padding='3 3 12 12')
mainframe.grid(column=0, row=0)
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(0, weight =1)

kaymakamlik = ttk.Entry(mainframe, width =30)
kaymakamlik.grid(column = 2, row = 0)

okul_adi = ttk.Entry(mainframe, width =30)
okul_adi.grid(column = 2, row = 1)

mudur_yard_adi = ttk.Entry(mainframe, width =30)
mudur_yard_adi.grid(column = 2, row = 2)

personel_adi_soyadi = ttk.Entry(mainframe, width =30)
personel_adi_soyadi.grid(column = 2, row = 3)

personel_gorev_brans = ttk.Entry(mainframe, width =30)
personel_gorev_brans.grid(column = 2, row = 4)

ttk.Label(mainframe, text ='KAYMAKAMLIK ADI').grid(column = 1, row = 0, sticky=W)
ttk.Label(mainframe, text ='OKULUN ADI').grid(column = 1, row = 1, sticky=W)
ttk.Label(mainframe, text ='MÜDÜR YARDIMCISININ ADI SOYADI').grid(column = 1, row = 2, sticky=W)
ttk.Label(mainframe, text ='PERSONELİN ADI SOYADI').grid(column = 1, row=3, sticky=W)
ttk.Label(mainframe, text ='PERSONELİN GÖREVİ/BRANŞI').grid(column = 1, row=4, sticky=W)

ttk.Label(mainframe, text ='MÜDÜR YARDIMCISI LİSTESİ').grid(column = 3, row=0)
ttk.Label(mainframe, text ='PERSONEL LİSTESİ').grid(column = 5, row=0)

ttk.Label(mainframe, text ='').grid(column = 5, row=31)
ttk.Label(mainframe, text ='').grid(column = 5, row=33)

ttk.Label(mainframe, text ='').grid(column = 3, row=31)
ttk.Label(mainframe, text ='').grid(column = 3, row=33)


liste2 = Listbox(mainframe,width=30)
liste2.grid(column=3, row=1,rowspan=30,  sticky=(N,S,E,W))
liste2.bind("<Double-Button-1>",bilgi_girişi2)

kaydirma2 = ttk.Scrollbar(mainframe, orient="vertical",command=liste2.yview)
kaydirma2.grid(column=4, row=1, rowspan=30,sticky='ns')

liste2.config(yscrollcommand=kaydirma2.set)
kaydirma2.config(command=liste2.yview)

liste = Listbox(mainframe,width=30)
liste.grid(column=5, row=1,rowspan=30,  sticky=(N,S,E,W))
liste.bind("<Double-Button-1>",bilgi_girişi)

kaydirma = ttk.Scrollbar(mainframe, orient="vertical",command=liste.yview)
kaydirma.grid(column=6, row=1, rowspan=30,sticky='ns')

liste.config(yscrollcommand=kaydirma.set)
kaydirma.config(command=liste.yview)

for i in sorted(os.listdir(),key=locale.strxfrm):
    if i.endswith('.sq3'):
        liste.insert(END,i[0:-4])

for i in sorted(os.listdir(),key=locale.strxfrm):
    if i.endswith('.sql3'):
        liste2.insert(END,i[0:-5])

vt1 = sqlite3.connect('okuladi.sql')
im1= vt1.cursor()
im1.execute("""CREATE TABLE IF NOT EXISTS imza(okuladi TEXT)""")
im1.execute("""SELECT * FROM  imza""")
rows = im1.fetchall()
data_str = ""
sf = "{}"
for row1 in rows:
    data_str += sf.format(row1[0])

okul_adi.insert(END,row1[0])

vt1.commit()

vt1 = sqlite3.connect('kaymakamlik.sq')
im1= vt1.cursor()
im1.execute("""CREATE TABLE IF NOT EXISTS imza(kaymakamlik TEXT)""")
im1.execute("""SELECT * FROM  imza""")
rows = im1.fetchall()
data_str = ""
sf = "{}"
for row in rows:
    data_str += sf.format(row[0])

kaymakamlik.insert(END,row[0])

vt1.commit()  

ttk.Button(mainframe, text='Personel Kaydet/Güncelle',command= kaydet_personel).grid(column=5, row=32)
ttk.Button(mainframe, text='Sil', command= sil_personel).grid(column=5, row=34)
ttk.Button(mainframe, text='İmza Sirküsü Ön İzleme', command = cikti).grid(column=2, row=32)

ttk.Button(mainframe, text='Müdür Yardımcısı Kaydet/Güncelle',command= kaydet_müd_yard).grid(column=3, row=32)
ttk.Button(mainframe, text='Sil', command= sil_müd_yard).grid(column=3, row=34)

kaymakamlik.focus()

root.mainloop()    
